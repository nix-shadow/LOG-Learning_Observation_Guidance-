#!/usr/bin/env python3
"""
LOG (Learning Observation Guidance) - Documentation Generator
Generates comprehensive, professionally formatted Word (.docx) documentation
for the LOG platform.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# --- Color Palette Constants ---
HEX_NAVY = "0A2540"      # Brand Blue
HEX_TEAL = "00B4D8"      # Brand Teal
HEX_AMBER = "FFB703"     # Brand Amber
HEX_DARK = "1E293B"      # Slate 800
HEX_GRAY_BG = "F8FAFC"   # Slate 50
HEX_BORDER = "CBD5E1"    # Slate 300
HEX_LIGHT_BLUE = "F0F9FF"# Sky 50
HEX_WHITE = "FFFFFF"
HEX_GREEN = "10B981"
HEX_RED = "EF4444"

COLOR_NAVY = RGBColor(10, 37, 64)
COLOR_TEAL = RGBColor(0, 180, 216)
COLOR_AMBER = RGBColor(255, 183, 3)
COLOR_DARK = RGBColor(30, 41, 59)
COLOR_GRAY = RGBColor(100, 116, 139)
COLOR_WHITE = RGBColor(255, 255, 255)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    """Set background color of a cell."""
    shading_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_xml)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style:
            b_el = OxmlElement(f'w:{border_name}')
            b_el.set(qn('w:val'), border_style.get('val', 'single'))
            b_el.set(qn('w:sz'), str(border_style.get('sz', 4)))
            b_el.set(qn('w:space'), '0')
            b_el.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(b_el)
        else:
            b_el = OxmlElement(f'w:{border_name}')
            b_el.set(qn('w:val'), 'none')
            tcBorders.append(b_el)
    tcPr.append(tcBorders)

def add_callout(doc, title, text, callout_type="info"):
    """Adds a stylish callout box with a colored left border and light background."""
    border_color = HEX_TEAL
    bg_color = HEX_LIGHT_BLUE
    title_color = COLOR_TEAL
    icon = "ℹ️ "

    if callout_type == "warning":
        border_color = HEX_AMBER
        bg_color = "FFFBEB"
        title_color = COLOR_AMBER
        icon = "⚠️ "
    elif callout_type == "rule":
        border_color = HEX_NAVY
        bg_color = HEX_GRAY_BG
        title_color = COLOR_NAVY
        icon = "📌 "
    elif callout_type == "success":
        border_color = HEX_GREEN
        bg_color = "F0FDF4"
        title_color = RGBColor(16, 185, 129)
        icon = "✅ "

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_shading(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    # Left border only
    set_cell_border(cell, 
                    left={'val': 'single', 'sz': 24, 'color': border_color},
                    top=None, bottom=None, right=None)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"{icon}{title}")
    run_t.font.name = "Arial"
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = title_color
    
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run_b = p2.add_run(text)
    run_b.font.name = "Arial"
    run_b.font.size = Pt(9.5)
    run_b.font.color.rgb = COLOR_DARK
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_heading_1(doc, text):
    h = doc.add_heading(level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_TEAL
    return h

def add_heading_3(doc, text):
    h = doc.add_heading(level=3)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK
    return h

def add_body_p(doc, text, bold_prefix=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_DARK
    return p

def add_bullet_item(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_DARK
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_DARK
    return p

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_shading(cell, "1E293B")  # Dark slate background
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    set_cell_border(cell, 
                    top={'val': 'single', 'sz': 4, 'color': '334155'},
                    bottom={'val': 'single', 'sz': 4, 'color': '334155'},
                    left={'val': 'single', 'sz': 4, 'color': '334155'},
                    right={'val': 'single', 'sz': 4, 'color': '334155'})
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(226, 232, 240)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_table(table, col_widths, headers, rows_data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        hdr_cells[i].width = col_widths[i]
        set_cell_shading(hdr_cells[i], HEX_NAVY)
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        set_cell_border(hdr_cells[i],
                        top={'val': 'single', 'sz': 6, 'color': HEX_NAVY},
                        bottom={'val': 'single', 'sz': 12, 'color': HEX_TEAL},
                        left={'val': 'single', 'sz': 4, 'color': '334155'},
                        right={'val': 'single', 'sz': 4, 'color': '334155'})
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = COLOR_WHITE

    # Data Rows
    for row_idx, r_data in enumerate(rows_data):
        row = table.add_row()
        r_cells = row.cells
        bg = HEX_GRAY_BG if row_idx % 2 == 1 else HEX_WHITE
        for col_idx, text in enumerate(r_data):
            r_cells[col_idx].text = str(text)
            r_cells[col_idx].width = col_widths[col_idx]
            set_cell_shading(r_cells[col_idx], bg)
            set_cell_margins(r_cells[col_idx], top=80, bottom=80, left=120, right=120)
            set_cell_border(r_cells[col_idx],
                            top={'val': 'single', 'sz': 4, 'color': HEX_BORDER},
                            bottom={'val': 'single', 'sz': 4, 'color': HEX_BORDER},
                            left={'val': 'single', 'sz': 4, 'color': HEX_BORDER},
                            right={'val': 'single', 'sz': 4, 'color': HEX_BORDER})
            p = r_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_DARK

def generate_full_documentation(output_path="LOG_Project_Documentation.docx"):
    print(f"Generating comprehensive documentation at: {output_path}")
    doc = docx.Document()

    # Configure Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Configure Header and Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("LOG - Learning Observation Guidance | System Architecture & Specification")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = COLOR_GRAY

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Confidential & Proprietary — LOG Platform Engineering Team")
        frun.font.name = "Arial"
        frun.font.size = Pt(8)
        frun.font.color.rgb = COLOR_GRAY

    # ==========================================
    # COVER / TITLE SECTION
    # ==========================================
    logo_path = os.path.abspath("frontend/public/assets/log-logo.png")
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(20)
        p_logo.paragraph_format.space_after = Pt(20)
        run_img = p_logo.add_run()
        run_img.add_picture(logo_path, width=Inches(3.2))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("LOG: Learning Observation Guidance")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(16)
    r_sub = p_sub.add_run("Master System Architecture, Technical Specification & Engineering Guide")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_TEAL

    # Document Metadata Box
    meta_tbl = doc.add_table(rows=1, cols=1)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_tbl.autofit = False
    m_cell = meta_tbl.rows[0].cells[0]
    m_cell.width = Inches(5.5)
    set_cell_shading(m_cell, HEX_GRAY_BG)
    set_cell_margins(m_cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(m_cell,
                    top={'val': 'single', 'sz': 6, 'color': HEX_BORDER},
                    bottom={'val': 'single', 'sz': 6, 'color': HEX_BORDER},
                    left={'val': 'single', 'sz': 12, 'color': HEX_NAVY},
                    right={'val': 'single', 'sz': 6, 'color': HEX_BORDER})

    mp = m_cell.paragraphs[0]
    mp.paragraph_format.space_before = Pt(0)
    mp.paragraph_format.space_after = Pt(2)
    
    def add_meta_line(p_obj, label, value):
        r1 = p_obj.add_run(f"{label}: ")
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_DARK
        r2 = p_obj.add_run(f"{value}\n")
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = COLOR_DARK

    add_meta_line(mp, "Document Version", "1.0.0 (Production Release)")
    add_meta_line(mp, "Platform Stack", "Next.js 14 App Router, TypeScript, Go (Gin), GORM, SQLite/PostgreSQL, IndexedDB")
    add_meta_line(mp, "Target Demographic", "Low-connectivity and edge learning environments (e.g. Rural Nepal)")
    add_meta_line(mp, "Offline Resilience", "100% Offline-capable PWA with IndexedDB Mutating Sync Queue")
    add_meta_line(mp, "Security & RBAC", "3-Tier RBAC (ADMIN, MODERATOR, STUDENT), JWT HMAC, Bcrypt, Gin Validation")
    add_meta_line(mp, "Document Status", "Official Architectural Specification & Engineering Reference")

    doc.add_page_break()

    # ==========================================
    # 1. EXECUTIVE SUMMARY & SYSTEM OVERVIEW
    # ==========================================
    add_heading_1(doc, "1. Executive Summary & System Overview")
    
    add_body_p(doc, "The LOG (Learning Observation Guidance) platform is an advanced, resilient educational technology ecosystem engineered specifically to eliminate educational disparities in low-connectivity and infrastructure-constrained regions (such as rural Nepal). While modern digital learning platforms increasingly depend on high-bandwidth, continuous cloud connectivity, LOG is architected from the ground up as an offline-first, edge-ready smart learning companion.")
    
    add_body_p(doc, "LOG bridges the gap between self-directed learning and intelligent pedagogical oversight. By continuously capturing fine-grained learner telemetry, the system constructs reflective observations, diagnoses mastery gaps, and delivers actionable, positive guidance—even when completely disconnected from the central server.")

    add_callout(doc, "Mission Statement", 
                "To provide accessible, high-performance, and supportive digital learning experiences to every student, irrespective of geographical remoteness, cellular coverage, or internet reliability.", 
                "info")

    add_heading_2(doc, "1.1 Key Value Propositions")
    add_bullet_item(doc, "Seamless, uninterrupted access to course catalogs, interactive lessons, and quizzes without requiring an active internet connection.", "Offline-First Resilience: ")
    add_bullet_item(doc, "Mutating user actions (progress updates, quiz submissions, activity completion) are stored in client-side IndexedDB and automatically synchronized upon network recovery.", "Bi-Directional Background Sync: ")
    add_bullet_item(doc, "Granular role enforcement guaranteeing appropriate access control for Principals/HODs (ADMIN), Teachers (MODERATOR), and Learners (STUDENT).", "Multi-Tier Role-Based Security: ")
    add_bullet_item(doc, "All guidance and observations follow strict positive phrasing protocols, cultivating learner confidence and resilience.", "Supportive Pedagogical Framing: ")
    add_bullet_item(doc, "Blazing-fast Go Gin backend coupled with Next.js 14 App Router and localized SQLite persistence ensuring instant response times.", "High Efficiency & Low Resource Footprint: ")

    # ==========================================
    # 2. CORE GUIDING PRINCIPLES & ENGINEERING CONSTRAINTS
    # ==========================================
    add_heading_1(doc, "2. Core Guiding Principles & Engineering Constraints")
    
    add_body_p(doc, "Every component, service, and data pipeline in the LOG platform is strictly bound by three non-negotiable architectural mandates:")

    add_callout(doc, "Rule 1: Low-Connectivity First",
                "LOG is purpose-built for environments with intermittent or non-existent internet. Developers and agents must NEVER bypass the offline caching layer (src/lib/api.ts). Direct fetch calls that fail to pass through IndexedDB caching and mutation queueing are strictly prohibited.",
                "rule")

    add_callout(doc, "Rule 2: Supportive & Constructive Language Only",
                "Observations, evaluations, and guidance messages must consistently employ positive, supportive phrasing. Negative phrasing (such as 'You failed', 'Poor score', or 'Incorrect attempt') is banned. Instead, construct affirmative recommendations such as 'This area could use more practice' or 'You are making steady progress.'",
                "rule")

    add_callout(doc, "Rule 3: Deterministic & Grounded Analytics",
                "All student metrics, mastery scores, streak statistics, and guidance recommendations must be directly computed from verified database records. External hallucinated generative LLM calls for grading or telemetry are strictly disallowed without verified data ground-truth.",
                "rule")

    # ==========================================
    # 3. THE 5-STAGE LOG PEDAGOGICAL CYCLE
    # ==========================================
    add_heading_1(doc, "3. The 5-Stage LOG Pedagogical Framework")
    
    add_body_p(doc, "At the conceptual core of the platform is the proprietary LOG Learning Cycle—a continuous feedback loop engineered to foster self-efficacy, deep mastery, and deliberate practice:")

    cycle_table = doc.add_table(rows=1, cols=4)
    format_table(cycle_table, 
                 [Inches(1.2), Inches(1.5), Inches(2.3), Inches(1.5)],
                 ["Stage", "Core Action", "System Function", "Outcome"],
                 [
                     ["1. Learn", "Interactive Modules", "Presents bite-sized concepts and interactive checks", "Acquisition of new theoretical knowledge"],
                     ["2. Observe", "Habit & Telemetry", "Captures response latency, quiz scores, and streak data", "Transparent reflection of learner habits"],
                     ["3. Understand", "Cognitive Diagnostic", "Maps topic mastery and pinpoints specific difficulty zones", "Awareness of strengths and growth opportunities"],
                     ["4. Guide", "Targeted Action", "Formulates concrete, actionable next steps", "Clear pathway to targeted remediation"],
                     ["5. Improve", "Deliberate Practice", "Executes targeted practice modules and review sessions", "Long-term knowledge retention and mastery"]
                 ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # 4. HIGH-LEVEL SYSTEM ARCHITECTURE
    # ==========================================
    add_heading_1(doc, "4. High-Level System Architecture")

    add_body_p(doc, "The LOG platform employs a distributed client-server architecture with an integrated offline storage and synchronization tier:")

    add_code_block(doc, 
"""+-------------------------------------------------------------------------+
|                         CLIENT LAYER (Next.js 14 PWA)                   |
|  +---------------------+  +----------------------+  +----------------+  |
|  |   Student Portal    |  |    Teacher Portal    |  | Admin Console  |  |
|  +---------------------+  +----------------------+  +----------------+  |
|                         |                          |                    |
|                         v                          v                    |
|  +-------------------------------------------------------------------+  |
|  |                 OFFLINE RESILIENCE INTERCEPTOR                    |  |
|  |                    (src/lib/api.ts fetchWithCache)                |  |
|  +-------------------------------------------------------------------+  |
|            |                                              |             |
|   [Offline / Network Err]                       [Online Network]        |
|            v                                              v             |
|  +-----------------------+                      +--------------------+  |
|  | IndexedDB: api-cache  |                      | REST API Request   |  |
|  | IndexedDB: sync-queue |                      | (Bearer JWT Token) |  |
|  +-----------------------+                      +--------------------+  |
+------------|----------------------------------------------|-------------+
             |                                              |
             | (Auto-flushed on 'online' event)             |
             +--------------------+                         |
                                  v                         v
+-------------------------------------------------------------------------+
|                         BACKEND LAYER (Go / Gin)                        |
|  +-------------------------------------------------------------------+  |
|  | Middleware: CORS, Security Headers, JWT HMAC Auth, Multi-Tier RBAC |  |
|  +-------------------------------------------------------------------+  |
|            |                          |                         |       |
|            v                          v                         v       |
|  +--------------------+      +--------------------+   +---------------+ |
|  | /api/auth/*        |      | /api/dashboard     |   | /api/admin/*  | |
|  | OTP, JWT, Session  |      | Progress, Guidance |   | Users, Roles  | |
|  +--------------------+      +--------------------+   +---------------+ |
|                                       |                                 |
|                                       v                                 |
|  +-------------------------------------------------------------------+  |
|  |                  GORM Object-Relational Mapping                   |  |
|  +-------------------------------------------------------------------+  |
|                                       |                                 |
|                                       v                                 |
|  +-------------------------------------------------------------------+  |
|  |            PERSISTENCE: SQLite (Local) / PostgreSQL (Prod)         |  |
|  +-------------------------------------------------------------------+  |
+-------------------------------------------------------------------------+""")

    add_heading_2(doc, "4.1 Technology Stack Rationale")
    tech_table = doc.add_table(rows=1, cols=3)
    format_table(tech_table,
                 [Inches(1.5), Inches(2.2), Inches(2.8)],
                 ["Component", "Technology", "Architectural Rationale"],
                 [
                     ["Frontend Framework", "Next.js 14 (App Router)", "Enables server/client component separation, instant routing, and high-performance bundle delivery."],
                     ["Offline PWA Layer", "next-pwa & Workbox", "Pre-caches static assets, application shell, and script bundles for instant offline boots."],
                     ["Local Data Storage", "IndexedDB (via 'idb')", "Provides asynchronous, high-capacity client-side storage for cached payloads and pending mutations."],
                     ["Styling & UI", "Tailwind CSS & Framer Motion", "Utility-first design system optimized for mobile responsiveness with 60fps micro-animations."],
                     ["Backend Engine", "Go 1.22+ (Gin Framework)", "High-concurrency, compiled binary with ultra-low memory footprint and sub-millisecond route handling."],
                     ["Database / ORM", "SQLite / GORM / PostgreSQL", "Zero-configuration auto-migrating local development SQLite database with seamless Postgres production parity."]
                 ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # 5. ADVANCED OFFLINE-FIRST SYNCING STRATEGY
    # ==========================================
    add_heading_1(doc, "5. Advanced Offline-First Synchronization Strategy")
    
    add_body_p(doc, "The offline engine located in src/lib/api.ts provides an abstraction over browser network operations. It intercepts all outgoing HTTP traffic and implements a dual-store IndexedDB architecture (database log-db v3, with a 24-hour cache TTL and exponential-backoff queue flushing):")

    add_heading_2(doc, "5.1 Dual-Store Architecture")
    add_bullet_item(doc, "Acts as an immutable Key-Value store where keys are API endpoints (e.g. '/dashboard') and values are serialized JSON response payloads. When online, successful GET requests automatically update this cache.", "api-cache Store: ")
    add_bullet_item(doc, "An auto-incrementing queue that records all offline mutating operations (POST, PUT, DELETE) including endpoint URL, HTTP method, headers, and request body payload.", "sync-queue Store: ")

    add_heading_2(doc, "5.2 Execution Flow Matrix")
    flow_table = doc.add_table(rows=1, cols=4)
    format_table(flow_table,
                 [Inches(1.2), Inches(1.2), Inches(1.8), Inches(2.3)],
                 ["Request Type", "Online Status", "Primary Action", "Fallback / Recovery Action"],
                 [
                     ["GET", "Online", "Executes fetch(); updates api-cache on 200 OK", "On network error: reads payload from api-cache"],
                     ["GET", "Offline", "Directly fetches cached JSON from api-cache", "Throws user-friendly offline notification if unvisited"],
                     ["POST/PUT/DELETE", "Online", "Executes fetch(); returns live response", "On network error: enqueues to sync-queue, returns 202"],
                     ["POST/PUT/DELETE", "Offline", "Enqueues payload in sync-queue; returns 202 Accepted", "Auto-flushes queue sequentially upon 'online' event"]
                 ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_2(doc, "5.3 Automatic Recovery & Queue Flushing")
    add_body_p(doc, "When the browser detects network restoration (via window.addEventListener('online')), the syncQueue() worker executes:")
    add_bullet_item(doc, "Retrieves all pending mutations in FIFO order from the 'sync-queue' store.")
    add_bullet_item(doc, "Re-issues each HTTP request to the backend with original headers and payloads.")
    add_bullet_item(doc, "Deletes successfully processed requests from the queue.")
    add_bullet_item(doc, "Displays a confirmation toast ('Offline changes synced successfully!') to the learner.")

    # ==========================================
    # 6. MULTI-TIER RBAC & SECURITY HARDENING
    # ==========================================
    add_heading_1(doc, "6. Multi-Tier RBAC & Security Hardening")
    
    add_body_p(doc, "The LOG platform implements a fortified 3-tier Role-Based Access Control model enforced both at the Gin API gateway and within the Next.js App Router navigation context:")

    rbac_table = doc.add_table(rows=1, cols=4)
    format_table(rbac_table,
                 [Inches(1.2), Inches(1.5), Inches(2.3), Inches(1.5)],
                 ["Role", "User Persona", "Permitted Access & Capabilities", "Route Guard"],
                 [
                     ["ADMIN", "Principal / HOD / Admin", "Full system oversight, user role reassignment, system-wide analytics, course & activity authoring", "/api/admin/*, /admin"],
                     ["MODERATOR", "Teacher / Educator", "Classroom roster inspection, learner streak tracking, grading, attention-flag reviews", "/api/moderator/*, /moderator"],
                     ["STUDENT", "Enrolled Learner", "Access to personal dashboard, catalog, interactive journey, observation metrics, guidance", "/api/dashboard, /learning, /observation, /guidance"]
                 ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_2(doc, "6.1 Cryptographic & Defensive Measures")
    add_bullet_item(doc, "All user passwords undergo salted Bcrypt hashing with a work factor of 14 (api/auth.go). Plaintext passwords are never persisted.", "Bcrypt Password Hashing: ")
    add_bullet_item(doc, "Stateless authentication via HMAC-SHA256 signed JSON Web Tokens (72-hour validity). Tokens are verified per request via the AuthMiddleware handler.", "HMAC-SHA256 JWT Tokens: ")
    add_bullet_item(doc, "All incoming JSON payloads are strictly validated using Gin binding struct tags (e.g. binding:\"required,min=10,max=15\"). Malformed requests are rejected with 400 Bad Request.", "Strict Schema Binding: ")
    add_bullet_item(doc, "Global middleware injects X-Content-Type-Options: nosniff, X-Frame-Options: DENY, and X-XSS-Protection: 1; mode=block.", "Hardened Security Headers: ")

    # ==========================================
    # 7. BACKEND API SPECIFICATION & REFERENCE
    # ==========================================
    add_heading_1(doc, "7. Backend API Specification & Reference")

    add_body_p(doc, "The Go backend runs on port 8080 by default. Below is the complete endpoint reference:")

    api_table = doc.add_table(rows=1, cols=5)
    format_table(api_table,
                 [Inches(0.8), Inches(1.8), Inches(1.1), Inches(1.3), Inches(1.5)],
                 ["Method", "Endpoint", "Access", "Description", "Key Params / Body"],
                 [
["POST", "/api/auth/request-otp", "Public", "Generates & sends 6-digit OTP", "{ phone: string }"],
                 ["POST", "/api/auth/verify-otp", "Public", "Validates OTP and issues JWT", "{ phone: string, otp: string }"],
                 ["POST", "/api/auth/logout", "Auth", "Revokes JWT via jti blocklist", "Bearer Token Header"],
                 ["GET", "/api/ping", "Public", "Health check probe", "None (Returns 200 pong)"],
                 ["GET", "/api/dashboard", "Student+", "Fetches learner profile, progress & guidance", "Bearer Token Header"],
                 ["GET", "/api/learning-journey", "Student+", "Lists all learning modules in sequence", "Bearer Token Header"],
                 ["GET", "/api/chart-data", "Student+", "Weekly performance & duration telemetry", "Bearer Token Header"],
                 ["GET", "/api/courses", "Student+", "Paginated course catalog", "page & limit query params"],
                 ["GET", "/api/activities/:id/modules", "Student+", "Bite-sized micro-module content", "URL :id activity"],
                 ["POST", "/api/activities/:id/complete", "Student+", "Transactional completion flow", "URL :id activity"],
                 ["POST", "/api/sync/bulk", "Student+", "Bulk offline .logsync upload", "{ version, timestamp, data[] }"],
                 ["GET", "/api/moderator/classes", "Moderator+", "Teacher class roster & student progress", "Bearer Token (MODERATOR/ADMIN)"],
                 ["GET", "/api/moderator/roster", "Moderator+", "Live roster w/ computed metrics", "Bearer Token (MODERATOR/ADMIN)"],
                 ["GET", "/api/admin/dashboard", "Admin", "System telemetry & user overview", "Bearer Token (ADMIN only)"],
                 ["GET", "/api/admin/users", "Admin", "Complete user registry listing", "Bearer Token (ADMIN only)"],
                 ["PUT", "/api/admin/users/:id/role", "Admin", "Updates target user role", "URL :id, { role: Role }"],
                 ["POST", "/api/admin/activities", "Admin", "Creates new learning module", "Strict CreateActivity DTO"]
                 ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # 8. DATABASE SCHEMA & DATA MODELS
    # ==========================================
    add_heading_1(doc, "8. Database Schema & Data Models")
    
    add_body_p(doc, "The database schema is auto-migrated on application startup using GORM. The relational entities are structured as follows:")

    add_heading_2(doc, "8.1 Entity Structure")
    
    add_bullet_item(doc, "Fields: ID (PK, string), Name (string), Email (uniqueIndex), Phone (uniqueIndex), PasswordHash (hidden), Role (STUDENT|MODERATOR|ADMIN), IsVerified (bool), CreatedAt, UpdatedAt, DeletedAt (soft delete).", "User: ")
    add_bullet_item(doc, "Fields: Phone (PK, string), OTP (string), ExpiresAt (time.Time). Enforces 5-minute time-to-live expiration.", "OTPRecord: ")
    add_bullet_item(doc, "Fields: ID (PK, string), Title (string), Description (string), Status (Completed|In progress|Pending), Topic (string), Order (int), ContentJSON (text), CreatedAt, DeletedAt.", "Activity: ")
    add_bullet_item(doc, "Fields: LearnerID (PK, string), TotalTopics (int), Completed (int), CurrentStreak (int), OverallScore (float64).", "Progress: ")
    add_bullet_item(doc, "Fields: ID (PK, string), LearnerID (FK, string), Category (strengths|areas needing improvement|consistency), Text (string), CreatedAt.", "Observation: ")
    add_bullet_item(doc, "Fields: ID (PK, string), LearnerID (FK, string), Text (string), Action (URL route), Type (next_step|practice|insight), CreatedAt.", "Guidance: ")

    # ==========================================
    # 9. FRONTEND PAGES & USER EXPERIENCE
    # ==========================================
    add_heading_1(doc, "9. Frontend Architecture & User Interface Design")

    add_body_p(doc, "The Next.js 14 frontend is structured with modular App Router pages and reusable UI components:")

    page_table = doc.add_table(rows=1, cols=3)
    format_table(page_table,
                 [Inches(1.8), Inches(1.5), Inches(3.2)],
                 ["Page Route", "Access Tier", "Core Features & Visual Components"],
                 [
                     ["/ (Landing)", "Public", "Hero section with LOG branding, value proposition, 5-stage LOG cycle interactive cards."],
                     ["/login", "Public", "2-step phone OTP verification, Google OAuth simulation, session hydration in AuthContext."],
                     ["/forgot-password", "Public", "Self-service password recovery flow with email reset validation."],
                     ["/dashboard", "Student+", "Streak indicator banner, circular daily goal meter, focus recommendations, activity timeline."],
                     ["/learning", "Student+", "Chronological module timeline with status indicators (Completed, In Progress, Locked)."],
                     ["/learning/[id]", "Student+", "Interactive multi-step quiz player with instant feedback, progress tracking, and supportive hints."],
                     ["/courses", "Student+", "24+ course catalog with multi-category filtering, full-text search, and difficulty badges."],
                     ["/observation", "Student+", "Mastery KPI cards, Recharts performance trend area charts, daily engagement bar charts."],
                     ["/guidance", "Student+", "Personalized actionable recommendations with direct action navigation buttons."],
                     ["/moderator", "Moderator+", "Teacher control center, classroom rosters, student progress meters, attention alerts."],
                     ["/admin", "Admin", "Principal control console, user registration table, live role editor, activity creator."]
                 ])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # 10. TESTING & VERIFICATION SUITE
    # ==========================================
    add_heading_1(doc, "10. Testing, Quality Assurance & Verification")

    add_body_p(doc, "The frontend offline synchronization layer is covered by comprehensive Jest unit tests (__tests__/api.test.ts):")

    add_bullet_item(doc, "Ensures that when online, fetchWithCache attempts network retrieval and populates the IndexedDB cache upon success.", "Test Case 1: Online Network Fetch: ")
    add_bullet_item(doc, "Verifies that when a network failure occurs, the application transparently falls back to cached data without crashing.", "Test Case 2: Network Error Fallback: ")
    add_bullet_item(doc, "Validates that when the browser window is offline, the fetch call is bypassed entirely and cache data is immediately returned.", "Test Case 3: Offline Direct Cache Access: ")

    add_code_block(doc, 
"""# Running Frontend Unit Tests
cd frontend
npx jest

# Test Results:
# PASS  __tests__/api.test.ts
# PASS  src/lib/syncExport.test.ts
# Test Suites: 2 passed, 2 total
# Tests:       11 passed, 11 total""")

    # ==========================================
    # 11. DEPLOYMENT & DEVOPS GUIDE
    # ==========================================
    add_heading_1(doc, "11. Installation, Setup & Deployment Guide")

    add_body_p(doc, "Follow the instructions below to run and deploy the LOG platform locally or in production:")

    add_heading_2(doc, "11.1 Backend Setup")
    add_code_block(doc, 
"""cd backend
go mod download
go build -o server main.go
./server
# Backend starts on http://localhost:8080 with auto-seeded database log.db""")

    add_heading_2(doc, "11.2 Frontend Setup")
    add_code_block(doc, 
"""cd frontend
npm install
# Create .env file: NEXT_PUBLIC_API_URL=http://localhost:8080/api
npm run build
npm start
# Frontend starts on http://localhost:3000""")

    add_heading_2(doc, "11.3 Docker Compose Deployment")
    add_code_block(doc, 
"""# Run full-stack environment with PostgreSQL
docker-compose up -d""")

    # ==========================================
    # 12. FUTURE ENHANCEMENTS & ROADMAP
    # ==========================================
    add_heading_1(doc, "12. Future Enhancements & Technical Roadmap")

    add_bullet_item(doc, "Complete Devanagari script support and localized pedagogical content for Nepali school curricula.", "Nepali Language Localization: ")
    add_bullet_item(doc, "Support for SMS and USSD-based quiz interaction for learners without smartphone access.", "SMS & Feature Phone Gateway: ")
    add_bullet_item(doc, "Local Wi-Fi mesh synchronization enabling students in a classroom to sync data directly with a teacher's device without internet.", "Peer-to-Peer Mesh Sync: ")
    add_bullet_item(doc, "A drag-and-drop course builder allowing rural educators to create customized interactive modules locally.", "Educator Authoring Studio: ")

    # Save Document
    doc.save(output_path)
    print(f"Documentation successfully generated at: {output_path}")

if __name__ == "__main__":
    generate_full_documentation()
